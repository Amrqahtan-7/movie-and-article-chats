import os
import io
import re
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import Optional

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_openai import ChatOpenAI
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate

from qdrant_client.http import models as qdrant_models

from pypdf import PdfReader
import docx

import chat_store
import user_store
import auth

# ---- Load environment variables ----
load_dotenv()
QDRANT_URL = os.getenv("QDRANT_URL")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
HF_API_KEY = os.getenv("HF_API_KEY")

UPLOAD_DIR = "data/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ---- Build the RAG pipeline once, when the server starts ----
print("Loading RAG pipeline...")

embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

vectorstore = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    url=QDRANT_URL,
    api_key=QDRANT_API_KEY,
    collection_name="movies_demo",
)

# Qdrant requires a payload index to exist before a field can be used in a
# filter — without this, every /chat request would fail with a 400 error
# the same way the cleanup script did on metadata.source.
try:
    vectorstore.client.create_payload_index(
        collection_name="movies_demo",
        field_name="metadata.uploaded_by",
        field_schema=qdrant_models.PayloadSchemaType.KEYWORD,
    )
    vectorstore.client.create_payload_index(
        collection_name="movies_demo",
        field_name="metadata.title",
        field_schema=qdrant_models.PayloadSchemaType.KEYWORD,
    )
    print("Ensured indexes on 'metadata.uploaded_by' and 'metadata.title'.")
except Exception as e:
    print(f"Index step: {e}")

llm = ChatOpenAI(
    model="meta-llama/Llama-3.1-8B-Instruct",
    openai_api_key=HF_API_KEY,
    openai_api_base="https://router.huggingface.co/v1",
)

system_prompt = (
    "Use the given context to answer the question. "
    "If you don't know the answer, say you don't know. "
    "Context: {context}"
)
prompt = ChatPromptTemplate.from_messages([
    ("system", system_prompt),
    ("human", "{input}"),
])

# NOTE: we no longer build a single global `retriever` / `qa_chain` here.
# A retriever built once at startup has no way to know which user is asking,
# so it would search everyone's uploads together. Instead, retrieval happens
# per-request inside /chat, scoped to the current user (see get_user_context).
question_answer_chain = create_stuff_documents_chain(llm, prompt)

from rag_graph import build_graph
chat_graph = build_graph(vectorstore, llm, question_answer_chain, lambda email: get_user_upload_dir(email))

from article_pipeline import (
    build_article_graph,
    get_user_articles_dir,
    verify_draft,
    revise_document,
    save_article,
    build_preview,
    default_status_text,
)
article_graph = build_article_graph(llm)

# In-memory store for articles currently being reviewed (not yet saved),
# keyed by user email. One pending article per user at a time. Resets on
# server restart — if you need that to survive restarts, this would move to
# a real store, same as chat_store/user_store.
pending_articles = {}

print("RAG pipeline ready.")

# ---- FastAPI app setup ----
app = FastAPI(title="Movie Chatbot API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

security = HTTPBearer()


def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> str:
    email = auth.decode_access_token(credentials.credentials)
    if not email:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    return email


# ---- Per-user helpers ----
def get_user_upload_dir(email: str) -> str:
    """Each user gets their own subfolder under data/uploads, so raw saved
    files never sit in one shared folder."""
    safe_id = email.lower().replace("@", "_at_").replace(".", "_")
    user_dir = os.path.join(UPLOAD_DIR, safe_id)
    os.makedirs(user_dir, exist_ok=True)
    return user_dir


def get_user_context(query: str, user_email: str, k: int = 3):
    """Search Qdrant restricted to:
       - this user's own uploaded chunks (metadata.uploaded_by == user_email), OR
       - the base movie dataset, which has no 'uploaded_by' owner at all.
    This is what actually prevents one user's question from pulling back
    another user's uploaded content — the restriction happens in the
    database query itself, not as an afterthought."""
    user_filter = qdrant_models.Filter(
        should=[
            qdrant_models.FieldCondition(
                key="metadata.uploaded_by",
                match=qdrant_models.MatchValue(value=user_email),
            ),
            qdrant_models.IsNullCondition(
                is_null=qdrant_models.PayloadField(key="metadata.uploaded_by")
            ),
        ]
    )
    return vectorstore.similarity_search(query, k=k, filter=user_filter)


# ---- Auth schemas ----
class AuthRequest(BaseModel):
    email: str
    password: str

class AuthResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    email: str


# ---- Auth endpoints ----
@app.post("/signup", response_model=AuthResponse)
def signup(request: AuthRequest):
    email = request.email.strip().lower()

    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Invalid email address")
    if len(request.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    if user_store.user_exists(email):
        raise HTTPException(status_code=400, detail="An account with this email already exists")

    password_hash = auth.hash_password(request.password)
    user_store.create_user(email, password_hash)

    token = auth.create_access_token(email)
    return AuthResponse(access_token=token, email=email)


@app.post("/login", response_model=AuthResponse)
def login(request: AuthRequest):
    email = request.email.strip().lower()
    user = user_store.get_user(email)

    if not user or not auth.verify_password(request.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Incorrect email or password")

    token = auth.create_access_token(email)
    return AuthResponse(access_token=token, email=email)


# ---- Basic endpoints ----
@app.get("/")
def root():
    return {"status": "Movie Chatbot API is running"}


# ---- Chat (session-aware + user-aware + upload-scoped) ----
class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None

class ChatResponse(BaseModel):
    answer: str
    session_id: str


@app.post("/chat", response_model=ChatResponse)
def chat(request: ChatRequest, current_user: str = Depends(get_current_user)):
    session_id = request.session_id or chat_store.create_session(current_user)

    chat_store.add_message(session_id, "user", request.message, current_user)

    result_state = chat_graph.invoke({
        "question": request.message,
        "user_email": current_user,
    })
    answer = result_state["final_answer"]

    chat_store.add_message(session_id, "assistant", answer, current_user)

    return ChatResponse(answer=answer, session_id=session_id)


# ---- Chat history endpoints (user-scoped) ----
@app.get("/sessions")
def list_sessions(current_user: str = Depends(get_current_user)):
    return chat_store.get_sessions(current_user)


@app.get("/sessions/{session_id}")
def get_session_messages(session_id: str, current_user: str = Depends(get_current_user)):
    return chat_store.get_messages(session_id, current_user)


@app.post("/sessions/new")
def new_session(current_user: str = Depends(get_current_user)):
    session_id = chat_store.create_session(current_user)
    return {"session_id": session_id}


@app.delete("/sessions/{session_id}")
def delete_session(session_id: str, current_user: str = Depends(get_current_user)):
    deleted = chat_store.delete_session(session_id, current_user)
    if not deleted:
        return {"error": "Session not found"}
    return {"status": "deleted", "session_id": session_id}


# ---- File upload: save to the user's own folder + extract text ----
def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in document.paragraphs)


@app.post("/upload")
async def upload_file(file: UploadFile = File(...), current_user: str = Depends(get_current_user)):
    file_bytes = await file.read()

    # os.path.basename strips any directory components a malicious filename
    # might contain (e.g. "../../etc/passwd"), so we only ever write inside
    # this user's own folder.
    safe_filename = os.path.basename(file.filename)
    filename_lower = safe_filename.lower()

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        return {"error": "Unsupported file type. Use PDF, DOCX, or TXT."}

    if not text.strip():
        return {"error": "No readable text found in this file."}

    user_dir = get_user_upload_dir(current_user)
    save_path = os.path.join(user_dir, safe_filename)
    base, ext = os.path.splitext(save_path)
    counter = 1
    while os.path.exists(save_path):
        save_path = f"{base}_{counter}{ext}"
        counter += 1

    with open(save_path, "wb") as f:
        f.write(file_bytes)

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    chunks = splitter.split_text(text)

    docs = [
        Document(
            page_content=chunk,
            metadata={"title": safe_filename, "source": "upload", "uploaded_by": current_user},
        )
        for chunk in chunks
    ]

    vectorstore.add_documents(docs)

    return {
        "status": "success",
        "filename": safe_filename,
        "saved_to": save_path,
        "chunks_added": len(chunks),
    }


# ---- Delete uploads: scoped strictly to the current user ----
@app.get("/uploads")
def list_my_uploads(current_user: str = Depends(get_current_user)):
    """List filenames currently sitting in this user's own upload folder."""
    user_dir = get_user_upload_dir(current_user)
    return {"files": os.listdir(user_dir)}


@app.delete("/uploads/{filename}")
def delete_one_upload(filename: str, current_user: str = Depends(get_current_user)):
    """Delete a single uploaded file — both the saved copy on disk and its
    chunks in Qdrant. The filter always includes uploaded_by == current_user,
    so this can never delete another user's file even if they guess the
    filename."""
    safe_filename = os.path.basename(filename)
    user_dir = get_user_upload_dir(current_user)
    file_path = os.path.join(user_dir, safe_filename)

    file_existed = os.path.exists(file_path)
    if file_existed:
        os.remove(file_path)

    vectorstore.client.delete(
        collection_name="movies_demo",
        points_selector=qdrant_models.FilterSelector(
            filter=qdrant_models.Filter(
                must=[
                    qdrant_models.FieldCondition(
                        key="metadata.title",
                        match=qdrant_models.MatchValue(value=safe_filename),
                    ),
                    qdrant_models.FieldCondition(
                        key="metadata.uploaded_by",
                        match=qdrant_models.MatchValue(value=current_user),
                    ),
                ]
            )
        ),
    )

    if not file_existed:
        return {"status": "not_found", "filename": safe_filename}
    return {"status": "deleted", "filename": safe_filename}


@app.delete("/uploads")
def delete_all_my_uploads(current_user: str = Depends(get_current_user)):
    """Delete ALL of the current user's uploaded files — disk copies and
    Qdrant chunks — but only theirs. Other users' uploads and the base
    dataset are untouched, since the filter always requires
    uploaded_by == current_user."""
    user_dir = get_user_upload_dir(current_user)
    deleted_files = os.listdir(user_dir)
    for f in deleted_files:
        os.remove(os.path.join(user_dir, f))

    vectorstore.client.delete(
        collection_name="movies_demo",
        points_selector=qdrant_models.FilterSelector(
            filter=qdrant_models.Filter(
                must=[
                    qdrant_models.FieldCondition(
                        key="metadata.uploaded_by",
                        match=qdrant_models.MatchValue(value=current_user),
                    )
                ]
            )
        ),
    )

    return {"status": "deleted", "files": deleted_files}


# ---- Multi-agent article pipeline: research -> write -> verify (loop) -> REVIEW -> save ----
class ArticleRequest(BaseModel):
    topic: str

class ReviseRequest(BaseModel):
    instruction: str

class ArticlePreviewResponse(BaseModel):
    status: str
    article: str        # markdown preview, not yet saved
    pending: bool = True

class ArticleSavedResponse(BaseModel):
    status: str
    article: str
    saved_path: str
    filename: str


@app.post("/generate-article", response_model=ArticlePreviewResponse)
def generate_article(request: ArticleRequest, current_user: str = Depends(get_current_user)):
    """Runs research -> write -> verify (looping as needed), then STOPS.
    Nothing is saved yet — the draft is held in memory for review."""
    result_state = article_graph.invoke({
        "topic": request.topic,
        "user_email": current_user,
        "attempts": 0,
    })

    status_text = default_status_text(result_state["verified"], result_state["attempts"])

    pending_articles[current_user] = {
        "topic": request.topic,      # stable, used only for the saved filename
        "title": request.topic,      # editable display title, starts equal to topic
        "status_text": status_text,  # editable status line, starts auto-generated
        "sources": result_state["sources"],
        "draft": result_state["draft"],
        "verified": result_state["verified"],   # drives the UI badge only, not the text
    }

    preview = build_preview(request.topic, status_text, result_state["draft"], result_state["sources"])
    return ArticlePreviewResponse(
        status="verified" if result_state["verified"] else "unverified",
        article=preview,
    )


@app.post("/article/revise", response_model=ArticlePreviewResponse)
def revise_article(request: ReviseRequest, current_user: str = Depends(get_current_user)):
    """Applies the instruction to the WHOLE letter — title, status line,
    body text, or the sources list — as one edit. The badge is
    re-verified against whatever sources remain so it stays accurate; the
    status_text LINE IN THE LETTER is left exactly as the user edited it
    (including removed), even if that no longer matches the badge — the
    badge is the source of truth for accuracy, the text is the user's own
    words to keep or change freely."""
    pending = pending_articles.get(current_user)
    if not pending:
        raise HTTPException(status_code=404, detail="No article is currently pending review. Generate one first.")

    instruction_lower = request.instruction.lower()
    wants_remove = "remove" in instruction_lower or "delete" in instruction_lower

    # Status/sources removal is handled directly here, BEFORE any LLM call,
    # for the common "remove/delete X" phrasing — instant and 100% reliable,
    # no model involved.
    if wants_remove and "status" in instruction_lower:
        pending["status_text"] = ""

    if wants_remove and "source" in instruction_lower:
        if "last" in instruction_lower and pending["sources"]:
            pending["sources"] = pending["sources"][:-1]
        elif "first" in instruction_lower and pending["sources"]:
            pending["sources"] = pending["sources"][1:]
        else:
            matched = False
            for s in list(pending["sources"]):
                if s["title"] and s["title"].lower() in instruction_lower:
                    pending["sources"].remove(s)
                    matched = True
                    break
            if not matched:
                # no specific one/qualifier named — treat as "remove all sources"
                pending["sources"] = []

    # For anything else (rewording, trimming to N sources, changing the
    # title, etc.), the LLM still runs. But rather than trying to predict
    # every possible phrasing that means "don't touch X" (a losing game —
    # tried a keyword blocklist before this and it kept missing real
    # phrasings), the server now enforces a hard rule after the fact: a
    # field is only allowed to change if its name was literally mentioned
    # in the instruction. If you didn't say "title", the title you get back
    # is guaranteed to be the exact one you had before, no matter what the
    # model decided to do with it — same for "status" and "source".
    revised = revise_document(
        llm, pending["title"], pending["status_text"], pending["draft"], pending["sources"], request.instruction
    )

    if "title" not in instruction_lower:
        revised["title"] = pending["title"]
    if "status" not in instruction_lower:
        revised["status_text"] = pending["status_text"]
    if "source" not in instruction_lower:
        revised["sources"] = pending["sources"]

    body_changed = revised["draft"].strip() != pending["draft"].strip()

    if body_changed:
        # The actual claims changed — re-check them against whatever sources
        # remain, since a body edit could introduce something unsupported.
        verify_result = verify_draft(llm, revised["draft"], revised["sources"])
        new_verified = verify_result["verified"]
    else:
        # Only the title/status/sources presentation changed, not the article
        # text itself. Re-running verification here would compare the SAME
        # already-confirmed body against a now-empty (or edited) sources
        # list and always fail — not because anything got less true, just
        # because there'd be nothing left to check it against. Keep the
        # verified state that was already earned by the actual content.
        new_verified = pending["verified"]

    pending["title"] = revised["title"] or pending["title"]
    pending["status_text"] = revised["status_text"]
    pending["draft"] = revised["draft"]
    pending["sources"] = revised["sources"]
    pending["verified"] = new_verified
    pending_articles[current_user] = pending

    preview = build_preview(pending["title"], pending["status_text"], pending["draft"], pending["sources"])
    return ArticlePreviewResponse(
        status="verified" if new_verified else "unverified",
        article=preview,
    )


@app.post("/article/confirm", response_model=ArticleSavedResponse)
def confirm_article(current_user: str = Depends(get_current_user)):
    """The user approved the current pending draft — save it to disk now."""
    pending = pending_articles.get(current_user)
    if not pending:
        raise HTTPException(status_code=404, detail="No article is currently pending review. Generate one first.")

    result = save_article(
        pending["topic"], pending["title"], pending["status_text"], pending["verified"],
        current_user, pending["draft"], pending["sources"],
    )
    del pending_articles[current_user]

    # save_article() names the body "final_article"; the response model calls it
    # "article", same as the preview response — map it across explicitly.
    return ArticleSavedResponse(
        status="verified" if pending["verified"] else "unverified",
        article=result["final_article"],
        saved_path=result["saved_path"],
        filename=result["filename"],
    )


@app.post("/article/discard")
def discard_article(current_user: str = Depends(get_current_user)):
    """Throws away the current pending draft without saving it."""
    pending_articles.pop(current_user, None)
    return {"status": "discarded"}


@app.get("/articles")
def list_my_articles(current_user: str = Depends(get_current_user)):
    """List this user's own saved articles, newest first. Only ever reads
    from their own folder, same isolation pattern as uploads."""
    user_dir = get_user_articles_dir(current_user)
    files = [f for f in os.listdir(user_dir) if f.endswith(".md")]
    files.sort(key=lambda f: os.path.getmtime(os.path.join(user_dir, f)), reverse=True)

    articles = []
    for f in files:
        path = os.path.join(user_dir, f)
        with open(path, "r", encoding="utf-8") as fh:
            raw = fh.read()

        meta_match = re.search(r"<!--\s*verified:(true|false)\s*-->\n?", raw)
        status = "verified" if (meta_match and meta_match.group(1) == "true") else "unverified"
        content = raw[meta_match.end():] if meta_match else raw

        # first non-empty line is "# title"
        lines = [l for l in content.split("\n") if l.strip()]
        topic = lines[0].lstrip("# ").strip() if lines else f

        articles.append({
            "filename": f,
            "topic": topic,
            "status": status,
            "modified_at": os.path.getmtime(path),
        })

    return articles


@app.get("/articles/{filename}")
def get_one_article(filename: str, current_user: str = Depends(get_current_user)):
    """Fetch the full content of one of this user's own saved articles.
    os.path.basename strips any path traversal attempt, and the file is only
    ever read from this user's own folder."""
    safe_filename = os.path.basename(filename)
    user_dir = get_user_articles_dir(current_user)
    path = os.path.join(user_dir, safe_filename)

    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Article not found")

    with open(path, "r", encoding="utf-8") as f:
        raw = f.read()

    meta_match = re.search(r"<!--\s*verified:(true|false)\s*-->\n?", raw)
    status = "verified" if (meta_match and meta_match.group(1) == "true") else "unverified"
    content = raw[meta_match.end():] if meta_match else raw  # hidden comment stripped for display

    return {"filename": safe_filename, "article": content, "status": status, "saved_path": path}


@app.delete("/articles/{filename}")
def delete_one_article(filename: str, current_user: str = Depends(get_current_user)):
    safe_filename = os.path.basename(filename)
    user_dir = get_user_articles_dir(current_user)
    path = os.path.join(user_dir, safe_filename)

    if not os.path.exists(path):
        return {"status": "not_found", "filename": safe_filename}

    os.remove(path)
    return {"status": "deleted", "filename": safe_filename}